import streamlit as st
import pandas as pd
import requests
import re
import time
import io
from urllib.parse import quote
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="DOI Paper Sorter", layout="wide")

st.title("📚 Paper DOI Extractor & Chronological Sorter")
st.write("Upload your spreadsheet, extract DOIs, fetch publication dates, and download a sorted Word document.")

# --- 1. User Inputs ---
col1, col2 = st.columns([2, 1])
with col1:
    uploaded_file = st.file_uploader("Upload an Excel (.xlsx) or CSV (.csv) file", type=['csv', 'xlsx'])
with col2:
    # Crossref requires an email for the Polite Pool. 
    # Providing a default credible-looking one, but allowing the user to change it.
    email_input = st.text_input(
        "Email for Crossref API (Polite Pool)", 
        value="researcher.data@academic-institution.edu",
        help="Providing your real email prevents Crossref from blocking your API requests."
    )

# --- 2. Helper Functions ---
def parse_date_parts(date_parts):
    """Safely parse Crossref date-parts array into a pandas datetime object."""
    if not date_parts:
        return None
    try:
        parts = date_parts[0]
        year = parts[0]
        month = parts[1] if len(parts) > 1 else 1
        day = parts[2] if len(parts) > 2 else 1
        return pd.to_datetime(f"{year}-{month}-{day}", errors='coerce')
    except Exception:
        return None

def parse_iso_date(date_value):
    """Parse an ISO date string into a pandas datetime object."""
    if not date_value:
        return None
    return pd.to_datetime(date_value, errors='coerce')


def empty_date_result():
    """Return an empty result row with stable column names."""
    return {
        'Published_Print': None,
        'Published_Online': None,
        'Issued_Any': None,
        'Best_Available_Date': None,
    }


def build_date_result(print_date=None, online_date=None, issued_date=None):
    """Build a consistent result payload and compute the earliest available date."""
    valid_dates = [d for d in [online_date, print_date, issued_date] if pd.notna(d)]
    best_date = min(valid_dates) if valid_dates else None
    return {
        'Published_Print': print_date,
        'Published_Online': online_date,
        'Issued_Any': issued_date,
        'Best_Available_Date': best_date,
    }


def has_any_date(date_result):
    """Check whether a result payload contains at least one usable date."""
    return any(pd.notna(date_result[column]) for column in date_result)


def clean_doi(doi):
    """Normalize DOI values coming from URLs or plain text cells."""
    return str(doi).replace('https://doi.org/', '').replace('http://doi.org/', '').strip()


def fetch_crossref_dates(doi, email):
    """Fetch publication dates from Crossref."""
    doi_clean = clean_doi(doi)
    headers = {'User-Agent': f'PaperSorterApp/1.0 (mailto:{email})'}
    url = f"https://api.crossref.org/works/{quote(doi_clean, safe='')}"

    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()
    data = response.json().get("message", {})

    return build_date_result(
        print_date=parse_date_parts(data.get("published-print", {}).get("date-parts")),
        online_date=parse_date_parts(data.get("published-online", {}).get("date-parts")),
        issued_date=parse_date_parts(data.get("issued", {}).get("date-parts")),
    )


def fetch_openalex_dates(doi, email):
    """Fallback date lookup using OpenAlex when Crossref does not return usable dates."""
    doi_clean = clean_doi(doi)
    headers = {'User-Agent': f'PaperSorterApp/1.0 (mailto:{email})'}
    url = f"https://api.openalex.org/works/https://doi.org/{quote(doi_clean, safe='')}"

    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()
    data = response.json()
    publication_date = parse_iso_date(data.get("publication_date"))

    return build_date_result(issued_date=publication_date)


def fetch_publication_dates(doi, email):
    """Fetch publication dates with Crossref first, then OpenAlex as fallback."""
    if pd.isna(doi) or not doi:
        return empty_date_result()

    try:
        crossref_result = fetch_crossref_dates(doi, email)
        if has_any_date(crossref_result):
            return crossref_result
    except requests.RequestException:
        pass

    try:
        return fetch_openalex_dates(doi, email)
    except requests.RequestException:
        return empty_date_result()
    finally:
        time.sleep(0.1)

def generate_word_doc(df):
    """Generate a sorted Word document in memory."""
    doc = Document()
    doc.add_heading('Sorted Research Papers', 0)
    
    # Sort by the Best Date (fallback to Jan 1, 2026 for completely missing dates)
    df_sorted = df.copy()
    df_sorted['Sort_Date'] = df_sorted['Best_Available_Date'].fillna(pd.to_datetime('2026-01-01'))
    df_sorted = df_sorted.sort_values(by='Sort_Date')
    
    # Create periods for headings (Month Year)
    df_sorted['Period'] = df_sorted['Sort_Date'].dt.to_period('M')
    unique_periods = df_sorted['Period'].dropna().unique()
    
    # Assume column index 1 has the citation text based on original file structure
    citation_col = df_sorted.columns[1] if len(df_sorted.columns) > 1 else df_sorted.columns[0]
    
    for period in sorted(unique_periods):
        doc.add_heading(period.strftime('%B %Y'), level=1)
        period_data = df_sorted[df_sorted['Period'] == period]
        
        for _, row in period_data.iterrows():
            citation = str(row[citation_col])
            doi = row['Extracted_DOI']
            best_date = row['Sort_Date'].strftime('%Y-%m-%d')
            
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{citation}").bold = False
            
            if pd.notna(doi):
                p.add_run(f"\nDOI: {doi}").italic = True
            
            date_run = p.add_run(f"\nDate Sorted: {best_date}")
            date_run.font.size = Pt(9)
            date_run.font.italic = True
            
    # Save to memory buffer
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. Main Application Logic ---
if uploaded_file is not None:
    # Load data
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file, encoding='latin-1')
        else:
            df = pd.read_excel(uploaded_file)
            
        st.success(f"Successfully loaded: {uploaded_file.name}")
        
        if st.button("Process DOIs and Fetch Dates", type="primary"):
            with st.spinner("Extracting DOIs and querying Crossref..."):
                # Extract DOIs
                doi_pattern = r'(10\.\d{4,9}/[-._;()/:A-Z0-9]+)'
                
                # Search across all string columns to find DOIs safely
                df_strings = df.select_dtypes(include=['object', 'string'])
                extracted = df_strings.apply(lambda col: col.str.extract(doi_pattern, flags=re.IGNORECASE)[0])
                df['Extracted_DOI'] = extracted.bfill(axis=1).iloc[:, 0]
                
                # Setup Progress Bar
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Fetch Dates safely
                total_rows = len(df)
                results = []
                
                for i, row in df.iterrows():
                    doi = row['Extracted_DOI']
                    dates = fetch_publication_dates(doi, email_input)
                    results.append(dates)
                    
                    # Update progress
                    progress = (i + 1) / total_rows
                    progress_bar.progress(progress)
                    status_text.text(f"Processed {i + 1} of {total_rows} records...")
                
                # Combine results back to DataFrame
                date_cols = pd.DataFrame(results)
                df = pd.concat([df, date_cols], axis=1)
                
                st.success("API processing complete!")
                
                # Display Results
                st.subheader("Data Preview with Publication Dates")
                st.dataframe(df[['Extracted_DOI', 'Best_Available_Date', 'Published_Print', 'Published_Online', 'Issued_Any']].head(10))
                
                # --- 4. Export & Download ---
                st.divider()
                st.subheader("Downloads")
                
                col_dl1, col_dl2 = st.columns(2)
                
                # Word Document Download
                word_file = generate_word_doc(df)
                with col_dl1:
                    st.download_button(
                        label="📄 Download Sorted Word Document",
                        data=word_file,
                        file_name="Chronological_Papers.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # Updated CSV Download
                csv_data = df.to_csv(index=False).encode('utf-8')
                with col_dl2:
                    st.download_button(
                        label="📊 Download Updated Table (CSV)",
                        data=csv_data,
                        file_name="Papers_with_Dates.csv",
                        mime="text/csv"
                    )
                
    except Exception as e:
        st.error(f"Error reading file: {e}")

# Footer
st.divider()
st.caption("Created by Nuno S. Osório")
